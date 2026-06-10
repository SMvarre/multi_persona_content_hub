"""
Read docs/WEEKLY_REPORT.md and save docs/WEEKLY_REPORT.docx for Microsoft Word.
Run: .venv/Scripts/python.exe scripts/convert_report_to_docx.py
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MD = ROOT / "docs" / "WEEKLY_REPORT.md"
DEFAULT_DOCX = ROOT / "docs" / "WEEKLY_REPORT.docx"


def remove_bold_and_code(text):
    """Strip **bold** and `code` markers from one line."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def is_table_separator(line):
    """True for markdown table divider rows like |---|---|."""
    s = line.strip().replace("|", "").strip()
    if not s:
        return False
    for ch in s:
        if ch not in "-: ":
            return False
    return True


def read_table_rows(lines, start_index):
    """Collect markdown table lines starting at start_index."""
    rows = []
    i = start_index
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if not is_table_separator(line):
            cells = [remove_bold_and_code(c.strip()) for c in line.strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a simple grid table to the document."""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_num, row in enumerate(rows):
        for col_num in range(col_count):
            text = row[col_num] if col_num < len(row) else ""
            table.rows[row_num].cells[col_num].text = text
    doc.add_paragraph()


def convert_markdown_to_docx(md_path, docx_path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # markdown table
        if stripped.startswith("|"):
            rows, i = read_table_rows(lines, i)
            add_table_to_doc(doc, rows)
            continue

        if stripped == "---":
            doc.add_paragraph("—" * 50)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(remove_bold_and_code(stripped[2:]), level=0)
        elif stripped.startswith("## "):
            doc.add_heading(remove_bold_and_code(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(remove_bold_and_code(stripped[4:]), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(remove_bold_and_code(stripped[5:]), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(remove_bold_and_code(stripped[2:]), style="List Bullet")
        elif stripped and stripped[0].isdigit() and ". " in stripped[:5]:
            dot = stripped.find(". ")
            body = stripped[dot + 2 :]
            doc.add_paragraph(remove_bold_and_code(body), style="List Number")
        elif stripped:
            doc.add_paragraph(remove_bold_and_code(stripped))

        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print("Wrote", docx_path)


if __name__ == "__main__":
    md_file = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_MD
    out_file = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_DOCX
    convert_markdown_to_docx(md_file, out_file)
